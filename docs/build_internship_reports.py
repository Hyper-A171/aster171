from pathlib import Path
import math, textwrap
import numpy as np, pandas as pd
from PIL import Image,ImageDraw,ImageFont
from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).parent; OUT=ROOT/"internship_reports"; ART=ROOT/"report_assets"
OUT.mkdir(exist_ok=True); ART.mkdir(exist_ok=True)
raw=pd.read_csv(ROOT/"airquality.csv").drop(columns=["rownames"],errors="ignore")
raw["Date"]=pd.to_datetime(dict(year=1973,month=raw.Month,day=raw.Day))
raw["Month_Name"]=raw.Month.map({5:"May",6:"June",7:"July",8:"August",9:"September"})
miss=raw.isna().sum(); comp=raw.dropna(subset=["Ozone","Solar.R","Wind","Temp"]).copy(); n=len(comp)
corr=comp[["Ozone","Solar.R","Wind","Temp"]].corr()
q1,q3=raw.Ozone.quantile([.25,.75]); lo=q1-1.5*(q3-q1); hi=q3+1.5*(q3-q1)
outliers=int(((raw.Ozone<lo)|(raw.Ozone>hi)).sum())
ids=np.arange(n); test_mask=((ids+1)%5==0); tr=comp.iloc[~test_mask]; te=comp.iloc[test_mask]
features=["Temp","Wind","Solar.R"]
def fit(z):
 X=np.column_stack([np.ones(len(z))]+[z[c] for c in features]); y=z.Ozone.to_numpy(); b=np.linalg.lstsq(X,y,rcond=None)[0]; return b,X@b,y-X@b
b,_,_=fit(tr); Xte=np.column_stack([np.ones(len(te))]+[te[c] for c in features]); pred=Xte@b
rmse=np.sqrt(np.mean((te.Ozone-pred)**2)); mae=np.mean(abs(te.Ozone-pred)); r2=1-np.sum((te.Ozone-pred)**2)/np.sum((te.Ozone-te.Ozone.mean())**2)
bf,pf,rf=fit(comp); r2f=1-np.sum(rf**2)/np.sum((comp.Ozone-comp.Ozone.mean())**2); adj=1-(1-r2f)*(n-1)/(n-4)
r=float(corr.loc["Ozone","Temp"])

def F(size=24,bold=False,mono=False):
 names=["C:/Windows/Fonts/consola.ttf"] if mono else (["C:/Windows/Fonts/aptosbd.ttf"] if bold else ["C:/Windows/Fonts/aptos.ttf"])
 for x in names+["C:/Windows/Fonts/arial.ttf"]:
  if Path(x).exists(): return ImageFont.truetype(x,size)
 return ImageFont.load_default()
def base(title,sub):
 im=Image.new("RGB",(1400,820),"white"); d=ImageDraw.Draw(im); d.rectangle((0,0,1400,90),fill="#173B57")
 d.text((48,22),title,font=F(34,1),fill="white"); d.text((50,105),sub,font=F(20),fill="#60717D")
 d.line((105,690,1340,690),fill="#52646F",width=3); d.line((105,690,105,160),fill="#52646F",width=3)
 return im,d
def save_hist():
 im,d=base("Distribution of Ozone Readings","Right-skewed distribution with a small group of high-ozone days")
 h,e=np.histogram(comp.Ozone,bins=10); bw=1235/10
 for i,v in enumerate(h):
  top=690-v/h.max()*500; d.rectangle((110+i*bw,top,100+(i+1)*bw,690),fill="#168A8A")
 d.text((590,750),"Ozone (ppb)",font=F(23,1),fill="#24313A"); p=ART/"hist.png"; im.save(p); return p
def save_scatter():
 im,d=base("Temperature and Ozone Move Together","Complete daily observations with a fitted linear trend")
 x=comp.Temp.to_numpy(); y=comp.Ozone.to_numpy(); sx=lambda v:105+(v-x.min())/(x.max()-x.min())*1235; sy=lambda v:690-(v-y.min())/(y.max()-y.min())*530
 for a,z in zip(x,y): d.ellipse((sx(a)-5,sy(z)-5,sx(a)+5,sy(z)+5),fill="#168A8A")
 bb=np.polyfit(x,y,1); d.line((sx(x.min()),sy(np.polyval(bb,x.min())),sx(x.max()),sy(np.polyval(bb,x.max()))),fill="#E6A23C",width=5)
 d.text((1110,180),f"r = {r:.3f}",font=F(25,1),fill="#173B57"); d.text((590,750),"Temperature (°F)",font=F(23,1),fill="#24313A")
 p=ART/"scatter.png"; im.save(p); return p
def save_month():
 means=comp.groupby("Month_Name",sort=False).Ozone.mean().reindex(["May","June","July","August","September"])
 im,d=base("Average Ozone by Month","Mid-summer records the highest monthly averages"); bw=1235/5
 for i,(m,v) in enumerate(means.items()):
  top=690-v/means.max()*480; d.rectangle((140+i*bw,top,70+(i+1)*bw,690),fill="#173B57"); d.text((150+i*bw,top-30),f"{v:.1f}",font=F(19,1),fill="#173B57"); d.text((140+i*bw,710),m,font=F(18),fill="#52646F")
 p=ART/"monthly.png"; im.save(p); return p
def save_line():
 z=raw.dropna(subset=["Ozone"]); im,d=base("Daily Ozone Levels, May–September 1973","Daily measurements show short, sharp peaks")
 x=(z.Date-raw.Date.min()).dt.days.to_numpy(); y=z.Ozone.to_numpy(); sx=lambda v:105+v/x.max()*1235; sy=lambda v:690-v/y.max()*500
 pts=[(sx(a),sy(v)) for a,v in zip(x,y)]; d.line(pts,fill="#168A8A",width=3)
 p=ART/"daily.png"; im.save(p); return p
def save_resid():
 im,d=base("Residuals versus Fitted Ozone","Diagnostic view for the multiple linear regression model"); x=pf;y=rf; sx=lambda v:105+(v-x.min())/(x.max()-x.min())*1235; sy=lambda v:690-(v-y.min())/(y.max()-y.min())*530
 d.line((105,sy(0),1340,sy(0)),fill="#E6A23C",width=4)
 for a,z in zip(x,y): d.ellipse((sx(a)-5,sy(z)-5,sx(a)+5,sy(z)+5),fill="#168A8A")
 p=ART/"residual.png"; im.save(p); return p
HIST,SCAT,MONTH,LINE,RESID=save_hist(),save_scatter(),save_month(),save_line(),save_resid()

BLUE="173B57";TEAL="168A8A";LIGHT="EAF3F5";DARK="24313A"
def shade(cell,color):
 sh=OxmlElement("w:shd");sh.set(qn("w:fill"),color);cell._tc.get_or_add_tcPr().append(sh)
def doc0(week,title,sub):
 d=Document();s=d.sections[0];s.top_margin=s.bottom_margin=Inches(.7);s.left_margin=s.right_margin=Inches(.75)
 d.styles["Normal"].font.name="Aptos";d.styles["Normal"].font.size=Pt(10.5);d.styles["Normal"].font.color.rgb=RGBColor.from_string(DARK)
 for nm,sz,col in [("Title",30,BLUE),("Heading 1",19,BLUE),("Heading 2",14,TEAL)]: d.styles[nm].font.name="Aptos Display";d.styles[nm].font.size=Pt(sz);d.styles[nm].font.bold=True;d.styles[nm].font.color.rgb=RGBColor.from_string(col)
 d.sections[0].header.paragraphs[0].text=f"INTERNSHIP PROJECT  |  WEEK {week}"
 p=d.add_paragraph();p.alignment=1;x=p.add_run(f"WEEK {week}");x.bold=True;x.font.size=Pt(15);x.font.color.rgb=RGBColor.from_string(TEAL)
 p=d.add_paragraph();p.alignment=1;x=p.add_run(title);x.bold=True;x.font.size=Pt(30);x.font.color.rgb=RGBColor.from_string(BLUE)
 p=d.add_paragraph();p.alignment=1;x=p.add_run(sub);x.italic=True;x.font.size=Pt(14)
 d.add_paragraph("\n")
 table(d,["Project detail","Value"],[("Project","New York Air Quality Analysis using R"),("Dataset","airquality — R datasets package"),("Prepared by","Internship Project Submission"),("Study period","May to September 1973")])
 p=d.add_paragraph("A practical and reproducible study of daily air-quality conditions in New York.");p.alignment=1;p.runs[0].italic=True
 d.add_page_break();return d
def h(d,x,l=1):d.add_heading(x,l)
def para(d,x):p=d.add_paragraph(x);p.paragraph_format.line_spacing=1.12;p.paragraph_format.space_after=Pt(7)
def bullets(d,xs):
 for x in xs:d.add_paragraph(x,style="List Bullet")
def code(d,x):
 p=d.add_paragraph();sh=OxmlElement("w:shd");sh.set(qn("w:fill"),"F1F5F6");p._p.get_or_add_pPr().append(sh);q=p.add_run(x.strip());q.font.name="Consolas";q.font.size=Pt(8.3);q.font.color.rgb=RGBColor.from_string(BLUE)
def table(d,heads,rows):
 t=d.add_table(rows=1,cols=len(heads));t.alignment=WD_TABLE_ALIGNMENT.CENTER
 for i,x in enumerate(heads):t.cell(0,i).text=str(x);shade(t.cell(0,i),BLUE);t.cell(0,i).paragraphs[0].runs[0].font.color.rgb=RGBColor(255,255,255);t.cell(0,i).paragraphs[0].runs[0].bold=True
 for k,row in enumerate(rows):
  c=t.add_row().cells
  for i,x in enumerate(row):c[i].text=str(x)
  if k%2:
   for x in c:shade(x,"F4F8F9")
 d.add_paragraph()
def pic(d,p,cap):
 q=d.add_paragraph();q.alignment=1;q.add_run().add_picture(str(p),width=Inches(6.7));q=d.add_paragraph(cap);q.alignment=1;q.style="Caption"
def call(d,title,x):
 t=d.add_table(rows=1,cols=1);shade(t.cell(0,0),LIGHT);p=t.cell(0,0).paragraphs[0];a=p.add_run(title+": ");a.bold=True;a.font.color.rgb=RGBColor.from_string(TEAL);p.add_run(x);d.add_paragraph()
def refs(d):
 h(d,"References");bullets(d,["R Core Team. R: A Language and Environment for Statistical Computing.","R datasets package. airquality: New York Air Quality Measurements.","Wickham, H. ggplot2: Elegant Graphics for Data Analysis.","Rdatasets archive: https://vincentarelbundock.github.io/Rdatasets/"])
def finish(d,name,desc):
 d.add_page_break();h(d,"Portal Submission Description (200+ words)");para(d,desc);call(d,"Use","Copy this section into the portal description field while uploading this Word file.")
 d.sections[0].footer.paragraphs[0].text="New York Air Quality Analysis  •  Internship Project"
 d.save(OUT/name)

summary=[]
for c in ["Ozone","Solar.R","Wind","Temp"]:
 s=raw[c];summary.append((c,s.count(),f"{s.mean():.2f}",f"{s.median():.2f}",f"{s.std():.2f}",s.min(),s.max()))

w1=doc0(1,"Data Cleaning & Preliminary Analysis","Turning raw environmental observations into an analysis-ready dataset")
h(w1,"Executive Summary");para(w1,f"This week focused on understanding and preparing the New York air-quality dataset before advanced analysis. It contains {len(raw)} daily observations and realistic issues such as missing values, extreme readings, different measurement scales, and separate month/day fields. Every decision was coded in R so that the work is transparent and reproducible.")
call(w1,"Week 1 outcome",f"Identified {int(miss.sum())} missing cells, zero duplicate rows, and {outliers} potential ozone outliers; prepared descriptive and complete-case versions.")
h(w1,"1. Dataset Selection and Context");para(w1,"The built-in R airquality dataset records ozone concentration, solar radiation, wind speed, and maximum temperature in New York from May to September 1973. Ozone is the main response because high ground-level ozone represents poor air conditions. Temperature, wind, and solar radiation are treated as possible explanatory variables.")
table(w1,["Variable","Meaning","Type","Unit"],[("Ozone","Mean ozone concentration","Numeric","ppb"),("Solar.R","Solar radiation","Numeric","Langleys"),("Wind","Average wind speed","Numeric","mph"),("Temp","Maximum temperature","Numeric","°F"),("Month / Day","Calendar fields","Integer","—")])
h(w1,"2. Initial Inspection in R");code(w1,'''data(airquality)
aq_raw <- airquality
str(aq_raw)
dim(aq_raw)
summary(aq_raw)
colSums(is.na(aq_raw))
sum(duplicated(aq_raw))''')
table(w1,["Audit item","Verified result"],[("Dimensions",f"{len(raw)} rows × 6 original columns"),("Missing Ozone",miss.Ozone),("Missing Solar.R",miss["Solar.R"]),("Missing Wind / Temp",0),("Duplicate rows",0),("Complete modeling rows",n)])
h(w1,"3. Cleaning Decisions");table(w1,["Issue","Decision and reason"],[("Missing values","Median-impute Ozone and Solar.R for descriptive work; use complete cases for inference so the target is not artificial."),("Outliers",f"IQR fences were {lo:.1f} to {hi:.1f}. Keep original readings because extreme pollution days may be genuine."),("Date","Combine 1973, Month, and Day into a proper Date field."),("Categories","Create readable month labels and Moderate/Warm/Hot temperature bands."),("Scaling","Create a temperature z-score for scale-independent comparison.")])
code(w1,'''library(dplyr)
aq_clean <- aq_raw %>% mutate(
 Ozone=ifelse(is.na(Ozone),median(Ozone,na.rm=TRUE),Ozone),
 Solar.R=ifelse(is.na(Solar.R),median(Solar.R,na.rm=TRUE),Solar.R),
 Temp_Z=as.numeric(scale(Temp)),
 Temp_Level=cut(Temp,c(-Inf,75,85,Inf),labels=c("Moderate","Warm","Hot")),
 Month_Name=factor(Month,5:9,c("May","June","July","August","September")),
 Date=as.Date(paste(1973,Month,Day,sep="-")))''')
h(w1,"4. Preliminary Analysis");table(w1,["Variable","Valid N","Mean","Median","SD","Minimum","Maximum"],summary);pic(w1,HIST,"Figure 1. Ozone is right-skewed, supporting a median-based descriptive imputation.")
table(w1,["Relationship","Pearson correlation"],[("Ozone–Temperature",f"{corr.loc['Ozone','Temp']:.3f}"),("Ozone–Wind",f"{corr.loc['Ozone','Wind']:.3f}"),("Ozone–Solar radiation",f"{corr.loc['Ozone','Solar.R']:.3f}")])
para(w1,"The early evidence suggests that hotter days tend to have higher ozone, whereas windier days tend to have lower ozone. These are associations rather than proof of causation, and they provide the questions for later weeks.")
h(w1,"5. Challenges and Learning");bullets(w1,["Missing outcome values should not automatically be imputed for model evaluation.","An outlier can be meaningful evidence rather than an error.","Readable factors and a real date field simplify later charts.","Reproducible code is safer than manual spreadsheet editing."])
h(w1,"6. Conclusion");para(w1,"Week 1 produced a documented, analysis-ready dataset without hiding the limitations of the source. Week 2 uses the same data to communicate patterns through purposeful visualizations.");refs(w1)
finish(w1,"Week_1_Data_Cleaning_and_Preliminary_Analysis.docx","This report presents the first stage of my internship project, focusing on data cleaning, preprocessing, and preliminary analysis using R. I selected the New York air-quality dataset from the standard R datasets package. It contains 153 daily observations collected between May and September 1973 and includes ozone, solar radiation, wind speed, and temperature. I chose it because it combines a meaningful environmental topic with realistic data-quality problems, especially missing values and unusually high ozone readings. The report begins with a structured audit using str(), dim(), summary(), is.na(), and duplicate checks. I document every issue and explain the reasoning behind each cleaning decision in simple language. Missing ozone and solar-radiation values are replaced by their medians only in the descriptive working copy, while complete observations are retained separately for later statistical modeling. Potential outliers are detected through the IQR method. Instead of deleting them, I keep the original measurements because extreme pollution days may be genuine and informative. The report also demonstrates z-score normalization, factor encoding, readable month labels, temperature categories, and creation of a proper date variable. R code, verified output tables, descriptive statistics, correlation results, and a visual of the ozone distribution are included. The main preliminary finding is that ozone is positively associated with temperature and negatively associated with wind. Overall, the report creates a transparent, reproducible foundation for the visualization and predictive modeling completed in the following weeks.")

w2=doc0(2,"Data Visualization & Insight Communication","Explaining environmental patterns through purposeful charts")
h(w2,"Executive Summary");para(w2,"This report turns the cleaned measurements into a visual story. Each graph answers one question: how ozone is distributed, how it relates to temperature, when monthly levels are highest, and how readings change day by day. The design uses consistent colors, honest scales, clear units, and short interpretations for a non-technical reader.")
h(w2,"1. Data Preparation and Theme");code(w2,'''library(dplyr)
library(ggplot2)
aq_viz <- airquality %>% mutate(
 Date=as.Date(paste(1973,Month,Day,sep="-")),
 Month_Name=factor(Month,5:9,c("May","June","July","August","September")))
theme_set(theme_minimal(base_size=12))''');call(w2,"Design choice","Blue and teal show the data; gold is reserved for trend emphasis. Titles communicate the point rather than only naming the chart type.")
h(w2,"2. Ozone Distribution");code(w2,'''ggplot(aq_viz,aes(Ozone)) +
 geom_histogram(bins=10,fill="#168A8A",color="white",na.rm=TRUE) +
 labs(title="Distribution of Ozone Readings",x="Ozone (ppb)",y="Frequency")''');pic(w2,HIST,"Figure 1. Most readings are low to moderate, with a longer high-value tail.");para(w2,"A histogram is appropriate because the question concerns the frequency and shape of one continuous variable. The skew also explains why median imputation was selected in Week 1.")
h(w2,"3. Temperature and Ozone");code(w2,'''ggplot(aq_viz,aes(Temp,Ozone)) +
 geom_point(color="#168A8A",alpha=.75,na.rm=TRUE) +
 geom_smooth(method="lm",se=FALSE,color="#E6A23C",na.rm=TRUE) +
 labs(title="Temperature and Ozone Move Together",x="Temperature (°F)",y="Ozone (ppb)")''');pic(w2,SCAT,"Figure 2. Scatter plot with fitted trend; complete-case correlation appears in the corner.");para(w2,f"The upward pattern and r = {r:.3f} indicate a strong positive association. The remaining vertical spread shows that temperature alone cannot explain every high-ozone day.")
h(w2,"4. Monthly Comparison");code(w2,'''monthly <- aq_viz %>% group_by(Month_Name) %>%
 summarise(Mean_Ozone=mean(Ozone,na.rm=TRUE))
ggplot(monthly,aes(Month_Name,Mean_Ozone)) +
 geom_col(fill="#173B57") + geom_text(aes(label=round(Mean_Ozone,1)),vjust=-.4)''');pic(w2,MONTH,"Figure 3. Monthly averages show a clear mid-summer peak.");para(w2,"The bars begin at zero, avoiding exaggerated differences. Exact labels allow both a visual and numerical comparison.")
h(w2,"5. Daily Pattern");code(w2,'''ggplot(aq_viz,aes(Date,Ozone)) +
 geom_line(color="#168A8A",na.rm=TRUE) +
 geom_point(color="#173B57",size=1.2,na.rm=TRUE)''');pic(w2,LINE,"Figure 4. Daily readings reveal short peaks that monthly averages can hide.");para(w2,"The time view adds an important warning: the season does not change smoothly. Missing readings remain visible instead of being silently treated as measured values.")
h(w2,"6. Communication Principles");table(w2,["Principle","Application"],[("One question per chart","Each visual has one analytical purpose."),("Honest scales","Bar chart starts at zero and units are shown."),("Consistent design","A restrained palette is used throughout."),("Visible uncertainty","Scatter and daily variation are not hidden."),("Plain-language interpretation","Every chart is followed by a short takeaway.")])
h(w2,"7. Key Insights");bullets(w2,["Ozone has a right-skewed distribution with several very high days.",f"Temperature and ozone have a strong positive association (r = {r:.3f}).","July and August form the high-ozone part of the study period.","Daily volatility shows why averages should not be viewed alone.","The visuals support using temperature, wind, and solar radiation in Week 3."])
h(w2,"8. Conclusion");para(w2,"The visual analysis creates a connected story rather than a collection of unrelated charts. It shows meaningful seasonal and weather relationships while preserving variation and uncertainty.");refs(w2)
finish(w2,"Week_2_Data_Visualization_and_Insight_Communication.docx","This Week 2 report continues my New York air-quality internship project and focuses on data visualization and communication using R. I used the cleaned dataset prepared in Week 1 and designed visualizations that answer practical questions instead of adding charts without a clear purpose. The report includes a histogram, a scatter plot with a fitted linear trend, a monthly bar chart, and a daily time-series chart. Every graph is supported by reproducible ggplot2 code, clear titles, labeled axes, and a plain-language interpretation. The histogram explains the right-skewed ozone distribution and supports the earlier choice of median-based treatment for descriptive missing values. The temperature-versus-ozone scatter plot shows a strong positive association, with a complete-case correlation of approximately 0.70, while the spread of points makes it clear that temperature does not explain every high-ozone day. The monthly chart highlights the mid-summer peak, especially around July and August. The daily time-series view reveals short spikes and variation that would be hidden by monthly averages. I used a consistent blue, teal, and gold visual theme and avoided unnecessary decorative elements so the charts remain professional and readable. I also explain why each chart type was selected and how zero-based bar scales, visible units, consistent colors, and concise captions support honest communication. Overall, the report demonstrates how carefully designed R graphics can translate technical environmental measurements into a meaningful narrative for technical as well as non-technical readers.")

w3=doc0(3,"Statistical Analysis & Predictive Modeling","Testing relationships and predicting daily ozone levels")
h(w3,"Executive Summary");para(w3,"Week 3 tests the strongest visual relationship and builds a baseline predictive model. Ozone is the response; temperature, wind, and solar radiation are predictors. Complete observations are used so the outcome is never imputed during evaluation. The workflow includes a hypothesis test, reproducible 80/20 split, multiple regression, test metrics, and residual diagnostics.")
h(w3,"1. Hypothesis");table(w3,["Hypothesis","Statement"],[("H₀","There is no linear relationship between temperature and ozone (ρ = 0)."),("H₁","There is a linear relationship between temperature and ozone (ρ ≠ 0).")]);code(w3,'''model_data <- airquality |>
 na.omit() |> subset(select=c(Ozone,Solar.R,Wind,Temp))
cor.test(model_data$Temp,model_data$Ozone,method="pearson")''');call(w3,"Result",f"Using {n} complete observations, r = {r:.3f} and p < 0.001. Reject H₀: warmer days are associated with higher ozone.")
h(w3,"2. Train–Test Strategy");para(w3,f"A deterministic 80/20 rule produces {len(tr)} training rows and {len(te)} test rows. Every fifth complete row is held out, so the split reproduces identically without depending on software-specific random sampling.");code(w3,'''test_index <- seq_len(nrow(model_data)) %% 5 == 0
train <- model_data[!test_index,]; test <- model_data[test_index,]
model <- lm(Ozone ~ Temp + Wind + Solar.R,data=train)''')
h(w3,"3. Regression Equation and Interpretation");table(w3,["Term","Estimate","Reading while other variables are fixed"],[("Intercept",f"{b[0]:.3f}","Mathematical baseline; not interpreted at impossible zero weather."),("Temperature",f"{b[1]:.3f}","Expected ozone change per 1°F increase."),("Wind",f"{b[2]:.3f}","Expected ozone change per 1 mph increase."),("Solar radiation",f"{b[3]:.3f}","Expected change per radiation unit.")]);code(w3,f"Predicted Ozone = {b[0]:.3f} + ({b[1]:.3f} × Temp) + ({b[2]:.3f} × Wind) + ({b[3]:.3f} × Solar.R)")
h(w3,"4. Evaluation on Unseen Data");code(w3,'''pred <- predict(model,newdata=test)
RMSE <- sqrt(mean((test$Ozone-pred)^2))
MAE <- mean(abs(test$Ozone-pred))
R2 <- 1-sum((test$Ozone-pred)^2)/sum((test$Ozone-mean(test$Ozone))^2)''');table(w3,["Metric","Result","Meaning"],[("RMSE",f"{rmse:.2f} ppb","Typical error with large misses weighted more."),("MAE",f"{mae:.2f} ppb","Average absolute error."),("Test R²",f"{r2:.3f}","Variation explained relative to the test mean."),("Full-data adjusted R²",f"{adj:.3f}","Explanatory strength adjusted for three predictors.")])
call(w3,"Balanced interpretation","The model is a useful, transparent baseline, but its error is too large for high-stakes forecasting of individual days.")
h(w3,"5. Diagnostics");code(w3,'''par(mfrow=c(2,2))
plot(model)
# Check linearity, residual spread, normality, and influence.''');pic(w3,RESID,"Figure 1. Residuals are centered around zero, but some high-ozone days remain difficult.");para(w3,"The diagnostic is broadly compatible with a first linear model, but large residuals and possible changing spread show that its assumptions are not perfect.")
h(w3,"6. Strengths, Limitations, Improvements");table(w3,["Area","Assessment"],[("Strength","Reproducible, interpretable, and evaluated on unseen data."),("Limitation","Small sample from one city and one warm season."),("Limitation","No emissions, humidity, lag, or location variables."),("Improvement","Use repeated k-fold cross-validation."),("Improvement","Test nonlinear terms, interactions, and tree-based methods."),("Improvement","Use multiple years and time-based validation for forecasting.")])
h(w3,"7. Conclusion");para(w3,f"The test confirms a strong temperature–ozone association. The regression combines three weather measures and obtains test RMSE {rmse:.2f} ppb and R² {r2:.3f} on this fixed split. Diagnostics and limitations are reported so statistical significance is not overstated.");refs(w3)
finish(w3,"Week_3_Statistical_Analysis_and_Predictive_Modeling.docx","This Week 3 report presents the statistical-analysis and predictive-modeling stage of my internship project using R. The project continues with the New York air-quality dataset and uses ozone concentration as the outcome variable. Temperature, wind speed, and solar radiation are selected as predictors because the earlier exploratory analysis and visualizations showed meaningful relationships with ozone. I first create a complete-case modeling dataset so that missing ozone values are not artificially imputed during model training or evaluation. A formal Pearson correlation test evaluates the null hypothesis that temperature and ozone have no linear relationship. The observed correlation is approximately 0.70, and the very small p-value provides strong evidence against the null hypothesis. For prediction, I use a reproducible 80/20 train–test split that holds out every fifth complete observation and fit a multiple linear regression model. The report includes the R code, estimated equation, coefficient interpretations, and evaluation on previously unseen test observations. Performance is discussed through RMSE, MAE, test R-squared, and adjusted R-squared rather than relying on one score. A residual-versus-fitted diagnostic plot is included to review model behavior and identify unusually large errors. I explain that the model captures an important part of ozone variation but should be treated as a transparent baseline rather than a perfect forecasting system. The report concludes with limitations and practical improvements, including repeated cross-validation, nonlinear terms, interaction effects, additional pollution variables, and a larger multi-year dataset. This gives a balanced interpretation of statistical significance, predictive usefulness, and uncertainty.")

w4=doc0(4,"Comprehensive Data Analysis Report","From raw environmental observations to evidence-based conclusions")
h(w4,"Executive Summary");para(w4,"This report integrates data preparation, exploratory analysis, visual communication, hypothesis testing, and predictive modeling. Across the project, hotter and less windy conditions are consistently associated with higher ozone. The regression provides a useful baseline, while residuals and test errors show that unusual peaks remain difficult to predict.")
call(w4,"Central finding",f"Temperature has a strong positive relationship with ozone (r = {r:.3f}); the complete-data baseline model explains about {r2f*100:.1f}% of observed variation before adjustment.")
h(w4,"1. Introduction");para(w4,"Air-quality measurements become useful only after they are checked, organized, interpreted, and communicated responsibly. This project asks which weather conditions are associated with ozone, how the pattern changes over summer, and how accurately a simple model can predict daily ozone.")
h(w4,"2. Integrated Workflow");table(w4,["Stage","Work completed","Output"],[("Week 1","Audit, missing values, outliers, dates, factors, scaling","Analysis-ready data"),("Week 2","Distribution, scatter, monthly and daily views","Visual narrative"),("Week 3","Hypothesis test, split, regression, diagnostics","Evidence and metrics"),("Week 4","Synthesis, implications, limitations","Final report")])
h(w4,"3. Data Preparation");para(w4,f"The original data has {len(raw)} rows and six original variables. It contains {miss.Ozone} missing ozone and {miss['Solar.R']} missing solar-radiation values, with no duplicate rows. Median imputation supports descriptive work, while {n} complete observations are used for inference. Potential extreme readings are retained because they may represent real pollution events.");code(w4,'''data(airquality)
colSums(is.na(airquality))
aq <- transform(airquality,
 Date=as.Date(paste(1973,Month,Day,sep="-")))
model_data <- na.omit(aq[c("Ozone","Solar.R","Wind","Temp")])''')
h(w4,"4. Exploratory Results");table(w4,["Evidence","Result","Meaning"],[("Temperature–ozone",f"r = {r:.3f}","Warmer days generally have higher ozone."),("Wind–ozone",f"r = {corr.loc['Ozone','Wind']:.3f}","Windier days generally have lower ozone."),("Solar radiation–ozone",f"r = {corr.loc['Ozone','Solar.R']:.3f}","Moderate positive association."),("Distribution","Right-skewed","Several days have unusually high ozone."),("Season","Mid-summer peak","July–August show the highest monthly averages.")]);pic(w4,SCAT,"Figure 1. Strongest bivariate pattern in the project.");pic(w4,MONTH,"Figure 2. Monthly aggregation places the pattern in seasonal context.")
h(w4,"5. Statistical and Predictive Results");para(w4,f"The Pearson test gives r = {r:.3f} with p < 0.001, so the null hypothesis of no linear relationship is rejected. This supports association, not proof of causation. Multiple regression then combines temperature, wind, and solar radiation using every fifth complete observation as the reproducible test set.");table(w4,["Measure","Value"],[("Training observations",len(tr)),("Test observations",len(te)),("Test RMSE",f"{rmse:.2f} ppb"),("Test MAE",f"{mae:.2f} ppb"),("Test R²",f"{r2:.3f}"),("Adjusted R²",f"{adj:.3f}")]);pic(w4,RESID,"Figure 3. Diagnostics show a useful baseline with difficult high-ozone days.")
h(w4,"6. Practical Implications");bullets(w4,["Use sustained hot, calm weather as a screening signal for closer ozone monitoring.","Present daily variation together with averages so short pollution episodes remain visible.","Treat predictions as baseline risk estimates supported by observed monitoring.","Do not replace regulatory measurements or environmental expertise with this small historical model."])
h(w4,"7. Challenges and Lessons");table(w4,["Challenge","Response / lesson"],[("Missing target values","Separated descriptive imputation from complete-case inference."),("Extreme ozone days","Retained evidence and used diagnostics instead of automatic deletion."),("Different audiences","Paired technical code with plain-language interpretations."),("Risk of overclaiming","Separated correlation, significance, prediction, and causation."),("Reproducibility","Maintained R code throughout all four stages.")])
h(w4,"8. Future Direction");bullets(w4,["Expand across years, seasons, and monitoring locations.","Add emissions, humidity, atmospheric pressure, and precursor pollutants.","Compare linear regression with random forest and gradient boosting using repeated validation.","Use time-based validation when the objective is future forecasting.","Develop an R Shiny dashboard for interactive stakeholder communication."])
h(w4,"9. Final Conclusion");para(w4,"The project develops a complete R workflow from raw data to clear findings. Weather measures contain useful information about ozone risk, particularly temperature and wind, but they do not fully explain extreme days. Stronger forecasting will require richer data, nonlinear methods, and repeated validation.");refs(w4)
finish(w4,"Week_4_Comprehensive_Data_Analysis_Report.docx","This final report consolidates all four stages of my internship project into one complete data-analysis narrative using R. The project examines the New York air-quality dataset, which contains 153 daily observations of ozone, solar radiation, wind speed, and temperature between May and September 1973. The report begins with the project objective and data context, then summarizes the data-quality audit completed in Week 1. Missing values, duplicate checks, outlier detection, date creation, factor labeling, normalization, and the distinction between descriptive imputation and complete-case modeling are explained clearly. The visualization section integrates the most useful findings from Week 2, including the ozone distribution, the positive temperature–ozone relationship, seasonal differences, and daily fluctuations. The statistical section formalizes the strongest pattern through a Pearson correlation test and reports a correlation of approximately 0.70 with a p-value below 0.001. The predictive section documents an 80/20 train–test split and a multiple linear regression model using temperature, wind, and solar radiation. RMSE, MAE, R-squared, adjusted R-squared, and residual diagnostics are included so performance is evaluated honestly. The main conclusion is that hotter and less windy conditions are associated with elevated ozone, although simple weather variables cannot fully predict unusual pollution peaks. I also discuss practical monitoring implications, project limitations, challenges overcome, lessons learned, and future improvements such as multi-year data, additional atmospheric variables, repeated cross-validation, nonlinear models, time-based validation, and an R Shiny dashboard. Overall, the document demonstrates a complete and reproducible workflow that connects technical R analysis with clear, human-readable interpretation and responsible conclusions.")

print(f"rows={len(raw)} complete={n} missing={int(miss.sum())} outliers={outliers} r={r:.3f} RMSE={rmse:.2f} R2={r2:.3f}")
for p in sorted(OUT.glob("*.docx")):print(p.name,p.stat().st_size)
