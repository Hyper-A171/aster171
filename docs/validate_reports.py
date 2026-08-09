from pathlib import Path
from docx import Document
import zipfile,re
for p in sorted(Path("docs/internship_reports").glob("*.docx")):
 d=Document(p)
 text=" ".join(x.text for x in d.paragraphs)
 desc=text.split("Portal Submission Description (200+ words)")[-1]
 images=sum(1 for x in d.part.rels.values() if "image" in x.reltype)
 ok=zipfile.ZipFile(p).testzip() is None
 print(p.name,"paragraphs",len(d.paragraphs),"tables",len(d.tables),"images",images,"desc_words",len(re.findall("[A-Za-z0-9'-]+",desc)),"zip_ok",ok)

